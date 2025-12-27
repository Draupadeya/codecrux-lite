# app.py
import os
import re
import json
import subprocess
import urllib.parse
import datetime
from flask import Flask, request, jsonify, render_template, session
from flask_cors import CORS
from pathlib import Path
from youtube_transcript_api import YouTubeTranscriptApi
import requests
import google.generativeai as genai
from dotenv import load_dotenv
from io import BytesIO
try:
    from docx import Document
except Exception:
    Document = None  # Fallback if python-docx isn't installed
import base64
import numpy as np
import cv2

# --- Load environment variables FIRST ---
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")

if not GEMINI_API_KEY:
    print("⚠ GEMINI_API_KEY not set; AI features will use fallbacks.")
else:
    print(f"✓ Loaded API Key: {GEMINI_API_KEY[:10]}...")

# --- Setup ---
app = Flask(__name__)
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})

# Configure Flask to serve static files from frontend directory
import os
FRONTEND_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'frontend'))
app.static_folder = FRONTEND_PATH
app.template_folder = FRONTEND_PATH
app.config['SESSION_TYPE'] = 'filesystem'

# Set secret key for session management (required for flask-session)
# Using environment variable or fallback for development
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')

if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
    except Exception as e:
        print(f"⚠ Failed to configure Gemini: {e}. Using fallbacks.")

TRANSCRIPTS_DIR = Path("./transcripts")
TRANSCRIPTS_DIR.mkdir(exist_ok=True)

# --- Helper Functions ---

def extract_youtube_id(url):
    """Extracts YouTube video ID from standard or short URLs."""
    parsed = urllib.parse.urlparse(url)
    if parsed.hostname in ['www.youtube.com', 'youtube.com']:
        q = urllib.parse.parse_qs(parsed.query)
        return q.get('v', [None])[0]
    elif parsed.hostname in ['youtu.be']:
        return parsed.path[1:]
    return None

def get_video_duration(video_id):
    """Get video duration in seconds using yt-dlp"""
    try:
        result = subprocess.run([
            "yt-dlp",
            "--print", "duration",
            f"https://www.youtube.com/watch?v={video_id}"
        ], capture_output=True, text=True, check=True)
        
        duration = float(result.stdout.strip())
        return int(duration)
    except Exception as e:
        print(f"Could not get video duration: {e}")
        return 3600

def get_transcript_with_timestamps(video_id):
    """Get transcript with timestamps."""
    transcript_json_path = TRANSCRIPTS_DIR / f"{video_id}_timestamps.json"
    
    if transcript_json_path.exists():
        with open(transcript_json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        with open(transcript_json_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_list, f)
        return transcript_list
    except Exception as e:
        print(f"YouTubeTranscriptApi failed: {e}")
        return None

def get_transcript_text(video_id):
    """Get plain text transcript (fallback)"""
    transcript_path = TRANSCRIPTS_DIR / f"{video_id}.txt"
    
    if transcript_path.exists():
        return transcript_path.read_text(encoding="utf-8")

    try:
        subprocess.run([
            "yt-dlp",
            "--write-auto-subs",
            "--sub-lang", "en",
            "--skip-download",
            "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
            f"https://www.youtube.com/watch?v={video_id}"
        ], check=True, capture_output=True, text=True)

        vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
        if vtt_files:
            vtt_file = vtt_files[0]
            lines = vtt_file.read_text(encoding="utf-8").splitlines()
            full_text = " ".join(
                re.sub(r'<[^>]+>', '', line) 
                for line in lines 
                if line.strip() 
                and "-->" not in line 
                and not line.strip().isdigit() 
                and "WEBVTT" not in line
                and not line.startswith("NOTE")
            )
            transcript_path.write_text(full_text, encoding="utf-8")
            vtt_file.unlink()
            return full_text
    except Exception as e:
        print(f"yt-dlp failed: {e}")
    
    raise Exception("Could not fetch transcript")

def generate_course_overview(full_transcript_text, course_title):
    """Generate 2-4 line overview of the entire course."""
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Use first 5000 characters for course overview
        context_text = full_transcript_text[:5000] if len(full_transcript_text) > 5000 else full_transcript_text
        
        prompt = f"""Analyze this video course transcript and provide a 2-4 sentence overview of what concepts and topics are taught.

Course Title: {course_title}

Transcript Sample:
{context_text}

Focus on:
- Main topics covered
- Key skills students will learn
- The learning journey from start to finish

Return only the overview, no extra text."""

        response = model.generate_content(prompt)
        overview = response.text.strip()
        overview = re.sub(r'\n+', ' ', overview)
        overview = re.sub(r'\s+', ' ', overview)
        
        return overview
        
    except Exception as e:
        print(f"Course overview generation failed: {e}")
        return f"A comprehensive {course_title} covering essential concepts and practical skills."

def generate_module_summary(module_text, module_num, start_time, end_time):
    """Generate summary for a specific time segment of the video."""
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Convert seconds to readable time
        start_min = int(start_time / 60)
        end_min = int(end_time / 60)
        
        # Use more text for better context
        context_text = module_text[:4000] if len(module_text) > 4000 else module_text
        
        prompt = f"""Summarize what is taught in this specific segment (minutes {start_min}-{end_min}) of a video course in 3-4 clear sentences.

Focus on:
- What specific topics/concepts are covered in THIS segment
- What the student will learn by watching THIS part
- Any key skills or knowledge gained

Transcript from this time segment:
{context_text}

Return only the summary, no extra text."""

        response = model.generate_content(prompt)
        summary = response.text.strip()
        summary = re.sub(r'\n+', ' ', summary)
        summary = re.sub(r'\s+', ' ', summary)
        
        return summary
        
    except Exception as e:
        print(f"Module summary generation failed: {e}")
        sentences = [s.strip() for s in module_text.split('.') if len(s.strip()) > 30]
        return '. '.join(sentences[:3]) + '.' if sentences else f"Content covering minutes {start_min} to {end_min} of the course."

def extract_key_points(module_text):
    """Extract 5 clear, distinct key points from the transcript."""
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        context_text = module_text[:4000] if len(module_text) > 4000 else module_text
        
        prompt = f"""Extract exactly 5 key learning points from this video transcript segment.
Each point should be:
- A single clear concept or topic
- 10-20 words maximum
- Actionable and specific

Transcript:
{context_text}

Return only a JSON array of strings, no other text:
["Point 1", "Point 2", "Point 3", "Point 4", "Point 5"]"""

        response = model.generate_content(prompt)
        points_text = response.text.strip()
        
        points_text = re.sub(r'^```json?\s*\n?', '', points_text, flags=re.MULTILINE)
        points_text = re.sub(r'\n?```\s*$', '', points_text, flags=re.MULTILINE)
        
        key_points = json.loads(points_text)
        
        if isinstance(key_points, list) and len(key_points) > 0:
            return key_points[:5]
        else:
            raise ValueError("Invalid format")
        
    except Exception as e:
        print(f"Key points extraction failed: {e}")
        sentences = [s.strip() for s in module_text.split('.') if 30 < len(s.strip()) < 150]
        return sentences[:5] if sentences else ["Understanding core concepts", "Practical implementation", "Best practices", "Common pitfalls", "Next steps"]

def split_transcript_with_timestamps(transcript_list, daily_study_minutes, video_duration, course_title):
    """Split transcript into modules with time-based AI summaries."""
    # 85% for video, 15% for quiz (applies to any daily study time)
    VIDEO_TIME_RATIO = 0.85
    QUIZ_TIME_RATIO = 0.15
    
    daily_study_seconds = daily_study_minutes * 60
    video_watch_seconds = int(daily_study_seconds * VIDEO_TIME_RATIO)
    quiz_seconds = int(daily_study_seconds * QUIZ_TIME_RATIO)
    
    # Calculate number of modules based on video watch time
    num_modules = max(1, int(video_duration / video_watch_seconds))
    seconds_per_module = video_duration / num_modules
    
    print(f"  📊 Time allocation per day:")
    print(f"     Total study time: {daily_study_minutes/60:.2f} hours")
    print(f"     Video watching: {video_watch_seconds/3600:.2f} hours (85%)")
    print(f"     Quiz time: {quiz_seconds/3600:.2f} hours (15%)")
    
    # Generate course overview from full transcript
    full_text = " ".join(seg['text'] for seg in transcript_list)
    course_overview = generate_course_overview(full_text, course_title)
    print(f"\n✓ Course Overview: {course_overview[:100]}...")
    
    modules = []
    
    for module_idx in range(num_modules):
        start_time = int(module_idx * seconds_per_module)
        end_time = int(min((module_idx + 1) * seconds_per_module, video_duration))
        
        # Get transcript segments for this time range
        module_segments = [
            seg for seg in transcript_list 
            if start_time <= seg['start'] < end_time
        ]
        
        module_text = " ".join(seg['text'] for seg in module_segments)
        
        print(f"\n  📝 Module {module_idx + 1}: {start_time}s - {end_time}s")
        print(f"     Generating time-based summary...")
        summary = generate_module_summary(module_text, module_idx + 1, start_time, end_time)
        print(f"     ✓ {summary[:80]}...")
        
        print(f"     Extracting key points...")
        key_points = extract_key_points(module_text)
        
        module_num = module_idx + 1
        # Calculate actual time spent on this module
        video_duration_hours = round(video_watch_seconds / 3600, 2)
        quiz_duration_hours = round(quiz_seconds / 3600, 2)
        total_duration_hours = round(daily_study_seconds / 3600, 2)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": video_duration_hours,
            "quizDuration": quiz_duration_hours,
            "totalDuration": total_duration_hours,
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {
                "description": summary,
                "keyPoints": key_points
            },
            "quiz": []
        })
    
    return modules, course_overview

def split_transcript_without_timestamps(transcript_text, daily_study_minutes, video_duration, course_title):
    """Fallback: Split by word count with AI summaries."""
    VIDEO_TIME_RATIO = 0.85
    QUIZ_TIME_RATIO = 0.15
    
    words = transcript_text.split()
    total_words = len(words)
    
    daily_study_seconds = daily_study_minutes * 60
    video_watch_seconds = int(daily_study_seconds * VIDEO_TIME_RATIO)
    quiz_seconds = int(daily_study_seconds * QUIZ_TIME_RATIO)
    
    num_modules = max(1, int(video_duration / video_watch_seconds))
    
    print(f"  📊 Time allocation per day:")
    print(f"     Total study time: {daily_study_minutes/60:.2f} hours")
    print(f"     Video watching: {video_watch_seconds/3600:.2f} hours (85%)")
    print(f"     Quiz time: {quiz_seconds/3600:.2f} hours (15%)")
    
    words_per_module = total_words // num_modules
    seconds_per_module = video_duration // num_modules
    
    # Generate course overview
    course_overview = generate_course_overview(transcript_text, course_title)
    print(f"\n✓ Course Overview: {course_overview[:100]}...")
    
    modules = []
    
    for i in range(num_modules):
        start_word = i * words_per_module
        end_word = min((i + 1) * words_per_module, total_words)
        
        module_text = " ".join(words[start_word:end_word])
        module_num = i + 1
        
        start_time = i * seconds_per_module
        end_time = min((i + 1) * seconds_per_module, video_duration)
        
        print(f"\n  📝 Module {module_num}: {start_time}s - {end_time}s")
        print(f"     Generating time-based summary...")
        summary = generate_module_summary(module_text, module_num, start_time, end_time)
        print(f"     ✓ {summary[:80]}...")
        
        print(f"     Extracting key points...")
        key_points = extract_key_points(module_text)
        
        video_duration_hours = round(video_watch_seconds / 3600, 2)
        quiz_duration_hours = round(quiz_seconds / 3600, 2)
        total_duration_hours = round(daily_study_seconds / 3600, 2)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": video_duration_hours,
            "quizDuration": quiz_duration_hours,
            "totalDuration": total_duration_hours,
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {
                "description": summary,
                "keyPoints": key_points
            },
            "quiz": []
        })
    
    return modules, course_overview

def generate_quiz(module_title, key_points):
    """Generate 5-question quiz using Gemini."""
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        key_points_text = "\n".join(f"- {point}" for point in key_points[:5])
        
        prompt = f"""Create exactly 5 multiple choice questions based on {module_title}.

Key points covered:
{key_points_text}

Return ONLY valid JSON (no markdown):
[
  {{
    "question": "Clear question?",
    "options": {{
      "A": "Option A",
      "B": "Option B",
      "C": "Option C",
      "D": "Option D"
    }},
    "correct_answer": "A"
  }}
]"""

        response = model.generate_content(prompt)
        quiz_text = response.text.strip()
        
        quiz_text = re.sub(r'^```json?\s*\n?', '', quiz_text, flags=re.MULTILINE)
        quiz_text = re.sub(r'\n?```\s*$', '', quiz_text, flags=re.MULTILINE)
        
        quiz = json.loads(quiz_text)
        return quiz[:5] if isinstance(quiz, list) else []
        
    except Exception as e:
        print(f"Quiz generation failed: {e}")
        return [{
            "question": f"What is a key concept from {module_title}?",
            "options": {
                "A": key_points[0] if len(key_points) > 0 else "Concept A",
                "B": key_points[1] if len(key_points) > 1 else "Concept B",
                "C": key_points[2] if len(key_points) > 2 else "Concept C",
                "D": "None of the above"
            },
            "correct_answer": "A"
        }]

# --- Routes ---

@app.route("/generate-plan", methods=["POST", "OPTIONS"])
def generate_plan():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        course_title = data.get("courseTitle", "Course")
        course_link = data.get("courseLink")
        daily_hours = float(data.get("dailyStudyHours", 1))
        
        video_id = extract_youtube_id(course_link)
        if not video_id:
            return jsonify({"error": "Invalid YouTube link"}), 400

        print(f"\n{'='*60}")
        print(f"📚 Processing: {course_title}")
        print(f"🎥 Video ID: {video_id}")
        print(f"⏱️  Daily study time: {daily_hours} hours")
        print(f"   ├─ Video watching: {daily_hours * 0.85:.2f} hours (85%)")
        print(f"   └─ Quiz time: {daily_hours * 0.15:.2f} hours (15%)")
        print(f"{'='*60}\n")
        
        video_duration = get_video_duration(video_id)
        print(f"✓ Video duration: {video_duration}s ({video_duration/60:.1f} minutes)")
        
        transcript_with_timestamps = get_transcript_with_timestamps(video_id)
        
        if transcript_with_timestamps:
            print(f"✓ Got transcript with timestamps: {len(transcript_with_timestamps)} segments")
            daily_plan, course_overview = split_transcript_with_timestamps(
                transcript_with_timestamps, 
                daily_hours * 60,
                video_duration,
                course_title
            )
        else:
            print("⚠ Using word-based splitting")
            transcript_text = get_transcript_text(video_id)
            print(f"✓ Got transcript: {len(transcript_text.split())} words")
            daily_plan, course_overview = split_transcript_without_timestamps(
                transcript_text,
                daily_hours * 60,
                video_duration,
                course_title
            )

        print(f"\n✓ Created {len(daily_plan)} modules\n")

        response_data = {
            "courseTitle": course_title,
            "courseDescription": course_overview,
            "videoID": video_id,
            "dailyPlan": daily_plan,
            "streak": 0,
            "progress": 0
        }
        
        print(f"✅ Plan generated successfully!\n")
        
        response = jsonify(response_data)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/generate-quiz", methods=["POST", "OPTIONS"])
def generate_quiz_route():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        module_title = data.get("moduleTitle")
        key_points = data.get("keyPoints", [])
        
        print(f"📝 Generating quiz for: {module_title}")
        quiz = generate_quiz(module_title, key_points)
        print(f"✓ Generated {len(quiz)} questions")
        
        response = jsonify(quiz)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"Error: {str(e)}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/get-motivation", methods=["POST", "OPTIONS"])
def get_motivation():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        module_title = data.get("moduleTitle")
        course_title = data.get("courseTitle")
        
        message = f"🎉 Excellent work! You've completed {module_title} in {course_title}. Keep building your knowledge!"
        
        response = jsonify({"message": message})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

# --- Integration with Django ---

@app.route("/courses", methods=["GET"])
def courses_portal():
    """
    Entry point for Django-redirected students.
    Accepts user info from Django and serves the StudyMate frontend.
    """
    try:
        # Get user info from query parameters (passed by Django)
        user = request.args.get('user', 'Guest')
        django_session = request.args.get('session_id', None)
        
        # Log the user entry
        print(f"\n✓ User '{user}' accessed courses portal from Django")
        if django_session:
            print(f"  Django Session: {django_session}")
        
        # Store user info in session for later use
        session['django_user'] = user
        session['django_session_id'] = django_session
        
        # Serve the StudyMate frontend (index.html)
        from flask import send_file
        frontend_file = os.path.join(FRONTEND_PATH, 'index.html')
        if os.path.exists(frontend_file):
            with open(frontend_file, 'r', encoding='utf-8') as f:
                content = f.read()
                # Inject user context into the HTML
                content = content.replace(
                    '<script src="script.js"></script>',
                    f'<script>window.djangoUser = "{user}";</script>\n    <script src="script.js"></script>'
                )
            return content, 200, {'Content-Type': 'text/html'}
        else:
            return jsonify({"error": "Frontend file not found"}), 404
    
    except Exception as e:
        print(f"❌ Error loading courses portal: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/", methods=["GET"])
def root():
    """Serve the home page or redirect to /courses"""
    # If accessed without Django context, redirect to home or show info
    try:
        user = session.get('django_user', 'Guest')
        if user != 'Guest':
            # User came from Django, serve the portal
            return courses_portal()
        else:
            # Direct access, serve index.html if it exists
            frontend_file = os.path.join(FRONTEND_PATH, 'index.html')
            if os.path.exists(frontend_file):
                with open(frontend_file, 'r', encoding='utf-8') as f:
                    return f.read(), 200, {'Content-Type': 'text/html'}
            else:
                return jsonify({
                    "message": "StudyMate API is running",
                    "docs": "Access /courses?user=username to begin",
                    "status": "ok"
                })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/style.css", methods=["GET"])
def serve_css():
    """Serve CSS file"""
    css_file = os.path.join(FRONTEND_PATH, 'style.css')
    if os.path.exists(css_file):
        with open(css_file, 'r', encoding='utf-8') as f:
            return f.read(), 200, {'Content-Type': 'text/css'}
    return "", 404, {'Content-Type': 'text/css'}

@app.route("/script.js", methods=["GET"])
def serve_js():
    """Serve JavaScript file"""
    js_file = os.path.join(FRONTEND_PATH, 'script.js')
    if os.path.exists(js_file):
        with open(js_file, 'r', encoding='utf-8') as f:
            return f.read(), 200, {'Content-Type': 'application/javascript'}
    return "", 404, {'Content-Type': 'application/javascript'}

@app.route("/favicon.ico", methods=["GET"])
def favicon():
    """Serve favicon (return empty 200 to prevent errors)"""
    return "", 200

@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "status": "healthy", 
        "message": "StudyMate API is running",
        "timestamp": str(datetime.datetime.now()),
        "integration": "Django @ http://localhost:8000",
        "cors": "Enabled for all origins"
    })

# =============================
# 🎥 Face Detection / Attention
# =============================

# Initialize a basic face detector (Haar cascade)
_FACE_DETECTOR = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')

def _decode_image_b64(image_b64: str):
    """Decode a base64 image (optionally data URI) into a BGR numpy array."""
    try:
        if image_b64.startswith('data:'):
            # Strip data URI header
            image_b64 = image_b64.split(',', 1)[1]
        raw = base64.b64decode(image_b64)
        np_arr = np.frombuffer(raw, dtype=np.uint8)
        img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
        return img
    except Exception:
        return None


def _detect_phone_like_regions(gray_img):
    """
    Naive phone detector based on rectangular, high-contrast regions.
    It is intentionally lightweight (no ML weights) and biased toward
    catching obvious handheld rectangles near the camera.
    """
    edges = cv2.Canny(gray_img, 50, 150)
    edges = cv2.dilate(edges, None, iterations=1)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    phone_boxes = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        area = w * h
        # Rough size and aspect ratio filters for a phone shape
        if area < 1200 or area > 45000:
            continue
        aspect = w / float(h)
        if aspect < 0.35 or aspect > 0.85:
            continue

        # Prefer near-rectangular contours
        perimeter = cv2.arcLength(cnt, True)
        approx = cv2.approxPolyDP(cnt, 0.04 * perimeter, True)
        if len(approx) < 4 or len(approx) > 8:
            continue

        fill_ratio = cv2.contourArea(cnt) / float(area)
        if fill_ratio < 0.45:
            continue

        phone_boxes.append((x, y, w, h))

    return {
        "phone_detected": len(phone_boxes) > 0,
        "phone_candidates": len(phone_boxes),
        "phone_boxes": phone_boxes,
    }

def _analyze_frame(img_bgr):
    """Return simple attention metrics from a single frame using heuristics."""
    h, w = img_bgr.shape[:2]
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    faces = _FACE_DETECTOR.detectMultiScale(gray, scaleFactor=1.2, minNeighbors=5, minSize=(60, 60))

    face_present = len(faces) > 0
    center_offset = 1.0
    faces_count = int(len(faces))

    if face_present:
        # Pick largest face
        x, y, fw, fh = max(faces, key=lambda f: f[2] * f[3])
        face_cx = x + fw / 2.0
        face_cy = y + fh / 2.0
        img_cx = w / 2.0
        img_cy = h / 2.0

        # Normalized center distance (0 = centered, ~1 = off-screen)
        dx = abs(face_cx - img_cx) / (w / 2.0)
        dy = abs(face_cy - img_cy) / (h / 2.0)
        center_offset = min(1.0, np.hypot(dx, dy))

    # Blur measure (variance of Laplacian); lower means blur/still
    blur_var = float(cv2.Laplacian(gray, cv2.CV_64F).var())
    brightness = float(gray.mean())

    phone_data = _detect_phone_like_regions(gray)
    multiple_faces = faces_count > 1

    # Heuristic attention score: face present and relatively centered
    attention_score = 0.0
    if face_present:
        attention_score = max(0.0, 1.0 - center_offset)

    distracted = (not face_present) or (center_offset > 0.35)
    # Heuristic boredom: face present, not distracted, but very low visual change (blur_var small)
    bored = face_present and (not distracted) and (blur_var < 30.0)

    return {
        "face_present": face_present,
        "multiple_faces": multiple_faces,
        "phone_detected": bool(phone_data["phone_detected"]),
        "phone_candidates": int(phone_data["phone_candidates"]),
        "attention_score": round(attention_score, 3),
        "distracted": bool(distracted),
        "bored": bool(bored),
        "metrics": {
            "center_offset": round(float(center_offset), 3),
            "blur_var": round(float(blur_var), 3),
            "brightness": round(float(brightness), 3),
            "faces_count": faces_count,
            "phone_boxes": phone_data["phone_boxes"],
            "frame_size": [int(w), int(h)]
        }
    }

@app.route("/analyze-face", methods=["POST", "OPTIONS"])
def analyze_face():
    """Analyze a webcam frame for attention/distraction/boredom using heuristics."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        image_b64 = data.get("frame")
        if not image_b64:
            response = jsonify({"error": "Missing 'frame' base64 image"})
            response.headers.add("Access-Control-Allow-Origin", "*")
            return response, 400

        img = _decode_image_b64(image_b64)
        if img is None:
            response = jsonify({"error": "Invalid image data"})
            response.headers.add("Access-Control-Allow-Origin", "*")
            return response, 400

        result = _analyze_frame(img)
        response = jsonify(result)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"analyze-face error: {e}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

# --- Notes (.docx) Generation ---
@app.route("/generate-notes-doc", methods=["POST", "OPTIONS"])
def generate_notes_doc():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        module_title = data.get("moduleTitle", "Module")
        summary = data.get("summary", "")
        key_points = data.get("keyPoints", [])

        if Document is None:
            # Provide a graceful fallback if python-docx is not available
            content = f"Notes for {module_title}\n\nSummary:\n{summary}\n\nKey Points:\n" + "\n".join(f"- {p}" for p in key_points[:5])
            blob = content.encode("utf-8")
            file_name = re.sub(r"[^\w\-]+", "_", module_title.strip()) or "module"
            file_name += ".txt"
            hex_blob = blob.hex()
        else:
            # Build a simple .docx file
            doc = Document()
            doc.add_heading(f"Notes: {module_title}", level=1)
            if summary:
                doc.add_paragraph("Summary:")
                doc.add_paragraph(summary)
            if key_points:
                doc.add_paragraph("Key Points:")
                for p in key_points[:5]:
                    doc.add_paragraph(p, style='List Bullet')

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            hex_blob = buf.read().hex()
            file_name = re.sub(r"[^\w\-]+", "_", module_title.strip()) or "module"
            file_name += ".docx"

        response = jsonify({
            "file_blob": hex_blob,
            "file_name": file_name
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"generate-notes-doc error: {e}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

# --- Practice Challenge Generation ---
@app.route("/generate-challenge", methods=["POST", "OPTIONS"])
def generate_challenge():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        module_title = data.get("moduleTitle", "Module")
        key_points = data.get("keyPoints", [])

        # Simple deterministic challenge without relying on the AI
        title = f"Build a Mini Page: {module_title}"
        question = (
            "Create a simple web page with a heading, a styled button, and a script that shows an alert when the button is clicked. "
            "Use semantic HTML, add basic CSS, and vanilla JS for interactivity."
        )
        starting_code = {
            "html": "<h1>Your Title</h1>\n<button id=\"actionBtn\">Click Me</button>",
            "css": "body { font-family: sans-serif; padding: 1rem; }\nbutton { padding: .5rem 1rem; }",
            "js": "document.getElementById('actionBtn').addEventListener('click', () => alert('Hello!'));"
        }
        solution = starting_code  # For now, solution equals starting code

        response = jsonify({
            "type": "web",
            "title": title,
            "question": question,
            "starting_code": starting_code,
            "solution": solution
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"generate-challenge error: {e}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

# --- AI Hint (deterministic helper) ---
@app.route("/get-hint", methods=["POST", "OPTIONS"])
def get_hint():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        challenge_question = data.get("challenge_question", "")
        user_code = data.get("user_code", {"html": "", "css": "", "js": ""})
        solution = data.get("solution", {"html": "", "css": "", "js": ""})
        try_count = int(data.get("try_count", 1))

        hints = []
        # Check HTML structure
        if "<h1" not in user_code.get("html", ""):
            hints.append("Add a main heading using <h1> to describe the page.")
        if "button" not in user_code.get("html", ""):
            hints.append("Include a <button id=\"actionBtn\"> element so you can attach a click handler.")

        # Check CSS basics
        if "font-family" not in user_code.get("css", ""):
            hints.append("Set a readable font in CSS, e.g., body { font-family: sans-serif; }.")
        if "padding" not in user_code.get("css", ""):
            hints.append("Add padding to the button for better click area.")

        # Check JS interaction
        js = user_code.get("js", "")
        if "addEventListener" not in js or "click" not in js:
            hints.append("Attach a click event listener to the button using addEventListener('click', ...).")

        if try_count <= 1 and not hints:
            hints.append("Try running your code to see the preview, then iterate on styling and interaction.")

        hint_text = "\n".join(hints) if hints else "Looks great! Consider improving accessibility and responsive styles."
        response = jsonify({"hint": hint_text})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"get-hint error: {e}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

# --- Session Management ---
@app.before_request
def init_session():
    """Initialize flask session if needed"""
    if 'django_user' not in session:
        session['django_user'] = 'Guest'

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎓 StudyMate API Server")
    print("="*60)
    print(f"Gemini Model: {GEMINI_MODEL}")
    print(f"Server: http://127.0.0.1:5000")
    print(f"CORS Enabled: All origins")
    print(f"Integration: Django @ http://localhost:8000")
    print("="*60 + "\n")
    app.run(debug=False, host='127.0.0.1', port=5000)