# app.py
import os
import re
import json
import subprocess
import urllib.parse
from flask import Flask, request, jsonify
from flask_cors import CORS
from pathlib import Path
from youtube_transcript_api import YouTubeTranscriptApi
import requests
from dotenv import load_dotenv
import cv2
import numpy as np
from deepface import DeepFace
import base64
from io import BytesIO
from PIL import Image

# Optional docx export for notes
try:
    from docx import Document
except ImportError:
    Document = None

# --- Load environment variables FIRST ---
load_dotenv()

MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
MISTRAL_MODEL = os.getenv("MISTRAL_MODEL", "mistral-small-latest")
MISTRAL_API_URL = os.getenv("MISTRAL_API_URL", "https://api.mistral.ai/v1/chat/completions")
ASMEBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")

if not MISTRAL_API_KEY:
    raise ValueError("MISTRAL_API_KEY not found in environment variables!")

print(f"✓ Loaded Mistral Key: {MISTRAL_API_KEY[:10]}...")

# --- Setup ---
app = Flask(__name__)
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})


TRANSCRIPTS_DIR = Path("./transcripts")
TRANSCRIPTS_DIR.mkdir(exist_ok=True)

# --- Helper Functions ---

def mistral_chat(prompt, temperature=0.25, max_tokens=400):
    """Call Mistral chat completion API and return text."""
    try:
        resp = requests.post(
            MISTRAL_API_URL,
            headers={
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": MISTRAL_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "max_tokens": max_tokens,
            },
            timeout=60,
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    except Exception as exc:
        print(f"Mistral call failed: {exc}")
        return ""

def analyze_learner_mood(frame_data):
    """
    Analyze learner's mood and engagement from webcam frame.
    Returns emotion analysis and engagement status.
    
    Args:
        frame_data: Base64 encoded image or image bytes
    
    Returns:
        dict: Analysis results including emotion, engagement level, and recommendations
    """
    try:
        # Decode base64 image if necessary
        if isinstance(frame_data, str):
            # Remove data URL prefix if present
            if ',' in frame_data:
                frame_data = frame_data.split(',')[1]
            img_bytes = base64.b64decode(frame_data)
            img = Image.open(BytesIO(img_bytes))
            img_array = np.array(img)
        else:
            img_array = frame_data
        
        # Convert RGB to BGR for OpenCV compatibility
        if len(img_array.shape) == 3 and img_array.shape[2] == 3:
            frame_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        else:
            frame_bgr = img_array
        
        # Perform emotion analysis using DeepFace
        analysis = DeepFace.analyze(
            frame_bgr, 
            actions=['emotion', 'age', 'gender'],
            enforce_detection=False,
            detector_backend='opencv'
        )
        
        # Handle both single face and multiple faces
        if isinstance(analysis, list):
            analysis = analysis[0]
        
        # Extract emotion data
        emotions = analysis.get('emotion', {})
        dominant_emotion = analysis.get('dominant_emotion', 'unknown')
        age = analysis.get('age', 'unknown')
        gender = analysis.get('dominant_gender', 'unknown')
        
        # Calculate engagement score based on emotions
        engagement_score = calculate_engagement_score(emotions, dominant_emotion)
        
        # Determine engagement status
        engagement_status = determine_engagement_status(engagement_score, dominant_emotion)
        
        # Generate recommendations
        recommendations = generate_engagement_recommendations(engagement_status, dominant_emotion)
        
        result = {
            'success': True,
            'dominant_emotion': dominant_emotion,
            'emotions': emotions,
            'engagement_score': engagement_score,
            'engagement_status': engagement_status,
            'recommendations': recommendations,
            'age': age,
            'gender': gender,
            'timestamp': None  # Will be set by caller if needed
        }
        
        return result
        
    except Exception as e:
        print(f"Mood analysis error: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'message': 'Could not analyze mood. Please ensure face is visible and well-lit.'
        }

def calculate_engagement_score(emotions, dominant_emotion):
    """
    Calculate engagement score (0-100) based on emotion probabilities.
    Higher scores indicate better engagement.
    """
    # Engaged emotions (positive indicators)
    engaged_emotions = ['happy', 'surprise', 'neutral']
    # Disengaged emotions (negative indicators)
    disengaged_emotions = ['sad', 'angry', 'fear', 'disgust']
    
    engaged_score = sum(emotions.get(e, 0) for e in engaged_emotions)
    disengaged_score = sum(emotions.get(e, 0) for e in disengaged_emotions)
    
    # Weight neutral slightly less than happy
    if 'neutral' in emotions:
        engaged_score = engaged_score - (emotions['neutral'] * 0.3)
    
    # Calculate final score (0-100)
    engagement_score = max(0, min(100, engaged_score - disengaged_score))
    
    return round(engagement_score, 2)

def determine_engagement_status(engagement_score, dominant_emotion):
    """
    Determine engagement status based on score and dominant emotion.
    """
    if engagement_score >= 70:
        if dominant_emotion in ['happy', 'surprise']:
            return 'highly_engaged'
        else:
            return 'engaged'
    elif engagement_score >= 40:
        if dominant_emotion == 'neutral':
            return 'neutral'
        else:
            return 'partially_engaged'
    elif engagement_score >= 20:
        if dominant_emotion in ['sad', 'fear']:
            return 'confused'
        else:
            return 'distracted'
    else:
        if dominant_emotion in ['sad', 'angry']:
            return 'frustrated'
        else:
            return 'bored'

def generate_engagement_recommendations(status, emotion):
    """
    Generate actionable recommendations based on engagement status.
    """
    recommendations_map = {
        'highly_engaged': [
            "Excellent focus! Keep up the great work!",
            "You're doing fantastic! Stay on track.",
            "Great engagement! Continue with this momentum."
        ],
        'engaged': [
            "Good focus! You're learning well.",
            "Nice work! Keep this concentration going.",
            "You're doing well! Stay focused."
        ],
        'neutral': [
            "Try to stay more engaged with the content.",
            "Consider taking notes to increase engagement.",
            "Stay active in your learning process."
        ],
        'partially_engaged': [
            "Your attention seems to be drifting. Refocus on the material.",
            "Take a deep breath and reconnect with the content.",
            "Consider taking a short break if needed."
        ],
        'confused': [
            "You seem confused. Try rewatching this section.",
            "Don't hesitate to pause and review unclear concepts.",
            "Consider taking notes or looking up additional resources."
        ],
        'distracted': [
            "You appear distracted. Minimize external interruptions.",
            "Remove distractions and refocus on the lecture.",
            "Try to eliminate background noise and focus."
        ],
        'frustrated': [
            "Take a short break if you're feeling frustrated.",
            "It's okay to pause. Come back when you're ready.",
            "Don't give up! Sometimes stepping away helps."
        ],
        'bored': [
            "You seem bored. Try changing your study environment.",
            "Consider taking a 5-minute break to refresh.",
            "Speed up the video or skip to more interesting sections."
        ]
    }
    
    return recommendations_map.get(status, ["Stay focused on your learning goals."])

def extract_youtube_id(url):
    """Extracts YouTube video ID from standard or short URLs."""
    parsed = urllib.parse.urlparse(url)
    if parsed.hostname in ['www.youtube.com', 'youtube.com']:
        q = urllib.parse.parse_qs(parsed.query)
        return q.get('v', [None])[0]
    elif parsed.hostname in ['youtu.be']:
        return parsed.path[1:]
    return None

def get_video_duration(video_id):
    """Get video duration in seconds using yt-dlp"""
    try:
        result = subprocess.run([
            "yt-dlp",
            "--print", "duration",
            f"https://www.youtube.com/watch?v={video_id}"
        ], capture_output=True, text=True, check=True)
        
        duration = float(result.stdout.strip())
        return int(duration)
    except Exception as e:
        print(f"Could not get video duration: {e}")
        return 3600

def get_transcript_with_timestamps(video_id):
    """Get transcript with timestamps."""
    transcript_json_path = TRANSCRIPTS_DIR / f"{video_id}_timestamps.json"
    
    if transcript_json_path.exists():
        with open(transcript_json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    try:
        transcript_list = None

        if hasattr(YouTubeTranscriptApi, "get_transcript"):
            try:
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
            except Exception as inner:
                print(f"YouTubeTranscriptApi.get_transcript failed: {inner}")

        if transcript_list is None and hasattr(YouTubeTranscriptApi, "list_transcripts"):
            try:
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id).find_transcript(['en']).fetch()
            except Exception as inner:
                print(f"YouTubeTranscriptApi.list_transcripts failed: {inner}")

        # Fallback to yt-dlp if YouTubeTranscriptApi is unavailable
        if transcript_list is None:
            print(f"Falling back to yt-dlp for {video_id}...")
            try:
                subprocess.run([
                    "yt-dlp", "--write-auto-subs", "--sub-lang", "en", "--skip-download",
                    "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
                    f"https://www.youtube.com/watch?v={video_id}"
                ], check=True, capture_output=True, text=True, timeout=60)
                
                vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
                if vtt_files:
                    vtt_file = vtt_files[0]
                    lines = vtt_file.read_text(encoding="utf-8").splitlines()
                    transcript_list = [
                        {"text": line, "start": 0, "duration": 0}
                        for line in lines
                        if line.strip() and "-->" not in line and not line.strip().isdigit() and "WEBVTT" not in line and not line.startswith("NOTE")
                    ]
                    vtt_file.unlink()
            except Exception as fallback_err:
                print(f"yt-dlp fallback failed: {fallback_err}")

        if transcript_list is None:
            print(f"No transcript available for {video_id}; using mock")
            return [{"text": f"Content for video {video_id}", "start": 0, "duration": 0}]

        with open(transcript_json_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_list, f)
        return transcript_list
    except Exception as e:
        print(f"get_transcript_with_timestamps failed: {e}")
        return [{"text": f"Content for video {video_id}", "start": 0, "duration": 0}]

def get_transcript_text(video_id):
    """Get plain text transcript (fallback)"""
    transcript_path = TRANSCRIPTS_DIR / f"{video_id}.txt"
    
    if transcript_path.exists():
        return transcript_path.read_text(encoding="utf-8")

    try:
        subprocess.run([
            "yt-dlp",
            "--write-auto-subs",
            "--sub-lang", "en",
            "--skip-download",
            "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
            f"https://www.youtube.com/watch?v={video_id}"
        ], check=True, capture_output=True, text=True)

        vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
        if vtt_files:
            vtt_file = vtt_files[0]
            lines = vtt_file.read_text(encoding="utf-8").splitlines()
            full_text = " ".join(
                re.sub(r'<[^>]+>', '', line) 
                for line in lines 
                if line.strip() 
                and "-->" not in line 
                and not line.strip().isdigit() 
                and "WEBVTT" not in line
                and not line.startswith("NOTE")
            )
            transcript_path.write_text(full_text, encoding="utf-8")
            vtt_file.unlink()
            return full_text
    except Exception as e:
        print(f"yt-dlp failed: {e}")
    
    raise Exception("Could not fetch transcript")

def generate_course_overview(full_transcript_text, course_title):
    """Generate 2-4 line overview of the entire course."""
    try:
        context_text = full_transcript_text[:5000] if len(full_transcript_text) > 5000 else full_transcript_text
        prompt = (
            "Analyze this video course transcript and provide a 2-4 sentence overview of what concepts and topics are taught.\n\n"
            f"Course Title: {course_title}\n\n"
            "Focus on main topics, key skills, and the learning journey. Return only the overview.\n\n"
            f"Transcript Sample:\n{context_text}"
        )
        overview = mistral_chat(prompt, temperature=0.2, max_tokens=220)
        overview = re.sub(r'\n+', ' ', overview)
        overview = re.sub(r'\s+', ' ', overview)
        return overview
    except Exception as e:
        print(f"Course overview generation failed: {e}")
        return f"A comprehensive {course_title} covering essential concepts and practical skills."

def generate_module_summary(module_text, module_num, start_time, end_time):
    """Generate summary for a specific time segment of the video."""
    try:
        # Convert seconds to readable time
        start_min = int(start_time / 60)
        end_min = int(end_time / 60)
        
        # Use more text for better context
        context_text = module_text[:4000] if len(module_text) > 4000 else module_text
        prompt = (
            f"Summarize what is taught in minutes {start_min}-{end_min} of a video course in 3-4 clear sentences.\n"
            "Focus on what is covered, what the student learns, and key skills.\n"
            f"Transcript from this time segment:\n{context_text}\n"
            "Return only the summary."
        )
        summary = mistral_chat(prompt, temperature=0.25, max_tokens=260)
        summary = re.sub(r'\n+', ' ', summary)
        summary = re.sub(r'\s+', ' ', summary)
        return summary
    except Exception as e:
        print(f"Module summary generation failed: {e}")
        sentences = [s.strip() for s in module_text.split('.') if len(s.strip()) > 30]
        return '. '.join(sentences[:3]) + '.' if sentences else f"Content covering minutes {start_min} to {end_min} of the course."

def extract_key_points(module_text):
    """Extract 5 clear, distinct key points from the transcript."""
    try:
        context_text = module_text[:4000] if len(module_text) > 4000 else module_text
        prompt = (
            "Extract exactly 5 key learning points from this video transcript segment.\n"
            "Each point should be a single clear concept, 10-20 words, actionable and specific.\n"
            f"Transcript:\n{context_text}\n"
            "Return only a JSON array of strings: [\"Point 1\", \"Point 2\", \"Point 3\", \"Point 4\", \"Point 5\"]"
        )
        points_text = mistral_chat(prompt, temperature=0.2, max_tokens=220)
        points_text = re.sub(r'^```json?\s*\n?', '', points_text, flags=re.MULTILINE)
        points_text = re.sub(r'\n?```\s*$', '', points_text, flags=re.MULTILINE)
        key_points = json.loads(points_text)
        if isinstance(key_points, list) and len(key_points) > 0:
            return key_points[:5]
        else:
            raise ValueError("Invalid format")
    except Exception as e:
        print(f"Key points extraction failed: {e}")
        sentences = [s.strip() for s in module_text.split('.') if 30 < len(s.strip()) < 150]
        return sentences[:5] if sentences else ["Understanding core concepts", "Practical implementation", "Best practices", "Common pitfalls", "Next steps"]

def split_transcript_with_timestamps(transcript_list, daily_study_minutes, video_duration, course_title):
    """Split transcript into modules with time-based AI summaries."""
    # 85% for video, 15% for quiz (applies to any daily study time)
    VIDEO_TIME_RATIO = 0.85
    QUIZ_TIME_RATIO = 0.15
    
    daily_study_seconds = daily_study_minutes * 60
    video_watch_seconds = int(daily_study_seconds * VIDEO_TIME_RATIO)
    quiz_seconds = int(daily_study_seconds * QUIZ_TIME_RATIO)
    
    # Calculate number of modules based on video watch time
    num_modules = max(1, int(video_duration / video_watch_seconds))
    seconds_per_module = video_duration / num_modules
    
    print(f"  📊 Time allocation per day:")
    print(f"     Total study time: {daily_study_minutes/60:.2f} hours")
    print(f"     Video watching: {video_watch_seconds/3600:.2f} hours (85%)")
    print(f"     Quiz time: {quiz_seconds/3600:.2f} hours (15%)")
    
    # Generate course overview from full transcript
    full_text = " ".join(seg['text'] for seg in transcript_list)
    course_overview = generate_course_overview(full_text, course_title)
    print(f"\n✓ Course Overview: {course_overview[:100]}...")
    
    modules = []
    
    for module_idx in range(num_modules):
        start_time = int(module_idx * seconds_per_module)
        end_time = int(min((module_idx + 1) * seconds_per_module, video_duration))
        
        # Get transcript segments for this time range
        module_segments = [
            seg for seg in transcript_list 
            if start_time <= seg['start'] < end_time
        ]
        
        module_text = " ".join(seg['text'] for seg in module_segments)
        
        print(f"\n  📝 Module {module_idx + 1}: {start_time}s - {end_time}s")
        print(f"     Generating time-based summary...")
        summary = generate_module_summary(module_text, module_idx + 1, start_time, end_time)
        print(f"     ✓ {summary[:80]}...")
        
        print(f"     Extracting key points...")
        key_points = extract_key_points(module_text)
        
        module_num = module_idx + 1
        # Calculate actual time spent on this module
        video_duration_hours = round(video_watch_seconds / 3600, 2)
        quiz_duration_hours = round(quiz_seconds / 3600, 2)
        total_duration_hours = round(daily_study_seconds / 3600, 2)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": video_duration_hours,
            "quizDuration": quiz_duration_hours,
            "totalDuration": total_duration_hours,
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {
                "description": summary,
                "keyPoints": key_points
            },
            "quiz": []
        })
    
    return modules, course_overview

def split_transcript_without_timestamps(transcript_text, daily_study_minutes, video_duration, course_title):
    """Fallback: Split by word count with AI summaries."""
    VIDEO_TIME_RATIO = 0.85
    QUIZ_TIME_RATIO = 0.15
    
    words = transcript_text.split()
    total_words = len(words)
    
    daily_study_seconds = daily_study_minutes * 60
    video_watch_seconds = int(daily_study_seconds * VIDEO_TIME_RATIO)
    quiz_seconds = int(daily_study_seconds * QUIZ_TIME_RATIO)
    
    num_modules = max(1, int(video_duration / video_watch_seconds))
    
    print(f"  📊 Time allocation per day:")
    print(f"     Total study time: {daily_study_minutes/60:.2f} hours")
    print(f"     Video watching: {video_watch_seconds/3600:.2f} hours (85%)")
    print(f"     Quiz time: {quiz_seconds/3600:.2f} hours (15%)")
    
    words_per_module = total_words // num_modules
    seconds_per_module = video_duration // num_modules
    
    # Generate course overview
    course_overview = generate_course_overview(transcript_text, course_title)
    print(f"\n✓ Course Overview: {course_overview[:100]}...")
    
    modules = []
    
    for i in range(num_modules):
        start_word = i * words_per_module
        end_word = min((i + 1) * words_per_module, total_words)
        
        module_text = " ".join(words[start_word:end_word])
        module_num = i + 1
        
        start_time = i * seconds_per_module
        end_time = min((i + 1) * seconds_per_module, video_duration)
        
        print(f"\n  📝 Module {module_num}: {start_time}s - {end_time}s")
        print(f"     Generating time-based summary...")
        summary = generate_module_summary(module_text, module_num, start_time, end_time)
        print(f"     ✓ {summary[:80]}...")
        
        print(f"     Extracting key points...")
        key_points = extract_key_points(module_text)
        
        video_duration_hours = round(video_watch_seconds / 3600, 2)
        quiz_duration_hours = round(quiz_seconds / 3600, 2)
        total_duration_hours = round(daily_study_seconds / 3600, 2)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": video_duration_hours,
            "quizDuration": quiz_duration_hours,
            "totalDuration": total_duration_hours,
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {
                "description": summary,
                "keyPoints": key_points
            },
            "quiz": []
        })
    
    return modules, course_overview

def generate_quiz(module_title, key_points):
        """Generate 5-question quiz using Mistral."""
        try:
                key_points_text = "\n".join(f"- {point}" for point in key_points[:5])
                prompt = f"""Create exactly 5 multiple choice questions based on {module_title}.

Key points covered:
{key_points_text}

Return ONLY valid JSON (no markdown):
[
    {{
        "question": "Clear question?",
        "options": {{
            "A": "Option A",
            "B": "Option B",
            "C": "Option C",
            "D": "Option D"
        }},
        "correct_answer": "A"
    }}
]"""

                quiz_text = mistral_chat(prompt, temperature=0.25, max_tokens=420)
                quiz_text = re.sub(r'^```json?\s*\n?', '', quiz_text, flags=re.MULTILINE)
                quiz_text = re.sub(r'\n?```\s*$', '', quiz_text, flags=re.MULTILINE)
                quiz = json.loads(quiz_text)
                return quiz[:5] if isinstance(quiz, list) else []
        
        except Exception as e:
                print(f"Quiz generation failed: {e}")
                return [{
                        "question": f"What is a key concept from {module_title}?",
                        "options": {
                                "A": key_points[0] if len(key_points) > 0 else "Concept A",
                                "B": key_points[1] if len(key_points) > 1 else "Concept B",
                                "C": key_points[2] if len(key_points) > 2 else "Concept C",
                                "D": "None of the above"
                        },
                        "correct_answer": "A"
                }]


# Lightweight face detector for attention tracking
_FACE_DETECTOR = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')


def decode_image_b64(image_b64):
    """Decode a base64 image (data URI allowed) into a BGR numpy array."""
    try:
        if not isinstance(image_b64, str):
            return None
        if image_b64.startswith('data:'):
            image_b64 = image_b64.split(',', 1)[1]
        raw = base64.b64decode(image_b64)
        np_arr = np.frombuffer(raw, dtype=np.uint8)
        return cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
    except Exception:
        return None


def analyze_frame_basic(img_bgr):
    """Return simple attention metrics from a single webcam frame."""
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    faces = _FACE_DETECTOR.detectMultiScale(gray, scaleFactor=1.2, minNeighbors=5, minSize=(60, 60))

    face_present = len(faces) > 0
    faces_count = int(len(faces))
    center_offset = 1.0

    if face_present:
        x, y, fw, fh = max(faces, key=lambda f: f[2] * f[3])
        face_cx = x + fw / 2.0
        face_cy = y + fh / 2.0
        img_cx = img_bgr.shape[1] / 2.0
        img_cy = img_bgr.shape[0] / 2.0
        dx = abs(face_cx - img_cx) / (img_bgr.shape[1] / 2.0)
        dy = abs(face_cy - img_cy) / (img_bgr.shape[0] / 2.0)
        center_offset = min(1.0, float(np.hypot(dx, dy)))

    blur_var = float(cv2.Laplacian(gray, cv2.CV_64F).var())
    brightness = float(gray.mean())

    attention_score = max(0.0, 1.0 - center_offset) if face_present else 0.0
    distracted = (not face_present) or center_offset > 0.35
    bored = face_present and (not distracted) and blur_var < 30.0

    return {
        "face_present": bool(face_present),
        "multiple_faces": faces_count > 1,
        "attention_score": round(attention_score, 3),
        "distracted": bool(distracted),
        "bored": bool(bored),
        "metrics": {
            "faces_count": faces_count,
            "center_offset": round(center_offset, 3),
            "blur_var": round(blur_var, 3),
            "brightness": round(brightness, 3)
        }
    }

# --- Routes ---

@app.route("/generate-plan", methods=["POST", "OPTIONS"])
def generate_plan():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        course_title = data.get("courseTitle", "Course")
        course_link = data.get("courseLink")
        daily_hours = float(data.get("dailyStudyHours", 1))
        
        video_id = extract_youtube_id(course_link)
        if not video_id:
            return jsonify({"error": "Invalid YouTube link"}), 400

        print(f"\n{'='*60}")
        print(f"📚 Processing: {course_title}")
        print(f"🎥 Video ID: {video_id}")
        print(f"⏱️  Daily study time: {daily_hours} hours")
        print(f"   ├─ Video watching: {daily_hours * 0.85:.2f} hours (85%)")
        print(f"   └─ Quiz time: {daily_hours * 0.15:.2f} hours (15%)")
        print(f"{'='*60}\n")
        
        video_duration = get_video_duration(video_id)
        print(f"✓ Video duration: {video_duration}s ({video_duration/60:.1f} minutes)")
        
        transcript_with_timestamps = get_transcript_with_timestamps(video_id)
        
        if transcript_with_timestamps:
            print(f"✓ Got transcript with timestamps: {len(transcript_with_timestamps)} segments")
            daily_plan, course_overview = split_transcript_with_timestamps(
                transcript_with_timestamps, 
                daily_hours * 60,
                video_duration,
                course_title
            )
        else:
            print("⚠ Using word-based splitting")
            transcript_text = get_transcript_text(video_id)
            print(f"✓ Got transcript: {len(transcript_text.split())} words")
            daily_plan, course_overview = split_transcript_without_timestamps(
                transcript_text,
                daily_hours * 60,
                video_duration,
                course_title
            )

        print(f"\n✓ Created {len(daily_plan)} modules\n")

        response_data = {
            "courseTitle": course_title,
            "courseDescription": course_overview,
            "videoID": video_id,
            "dailyPlan": daily_plan,
            "streak": 0,
            "progress": 0
        }
        
        print(f"✅ Plan generated successfully!\n")
        
        response = jsonify(response_data)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/generate-quiz", methods=["POST", "OPTIONS"])
def generate_quiz_route():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        module_title = data.get("moduleTitle")
        key_points = data.get("keyPoints", [])
        
        print(f"📝 Generating quiz for: {module_title}")
        quiz = generate_quiz(module_title, key_points)
        print(f"✓ Generated {len(quiz)} questions")
        
        response = jsonify(quiz)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"Error: {str(e)}")
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/get-motivation", methods=["POST", "OPTIONS"])
def get_motivation():
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
        
    try:
        data = request.get_json()
        module_title = data.get("moduleTitle")
        course_title = data.get("courseTitle")
        
        message = f"🎉 Excellent work! You've completed {module_title} in {course_title}. Keep building your knowledge!"
        
        response = jsonify({"message": message})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500


@app.route("/ask-question", methods=["POST", "OPTIONS"])
def ask_question():
    """AI tutor endpoint that answers a learner question using Mistral and transcript context."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        question = data.get("question")
        video_id = data.get("videoId")
        course_title = data.get("courseTitle", "this course")
        current_time = int(data.get("currentTime", 0))

        if not question or not video_id:
            return jsonify({"error": "Missing question or videoId"}), 400

        transcript_list = get_transcript_with_timestamps(video_id)
        if transcript_list:
            full_transcript = " ".join(entry['text'] for entry in transcript_list)
        else:
            full_transcript = get_transcript_text(video_id)

        if transcript_list:
            context_start = max(0, current_time - 120)
            context_end = current_time + 120
            context_entries = [
                entry for entry in transcript_list
                if context_start <= entry['start'] <= context_end
            ]
            context_text = " ".join(entry['text'] for entry in context_entries)
        else:
            video_duration = get_video_duration(video_id)
            chars_per_second = len(full_transcript) / max(video_duration, 1)
            context_start = max(0, int((current_time - 120) * chars_per_second))
            context_end = int((current_time + 120) * chars_per_second)
            context_text = full_transcript[context_start:context_end]

        prompt = f"""You are a helpful tutor for the course: \"{course_title}\".

    Student Question: {question}
    Current video time: {current_time}s

    Relevant transcript context:
    {context_text}

    Provide a clear, concise answer grounded in the context. If the context is thin, explain the concept with your own knowledge but note it may not have been covered yet."""
        answer = mistral_chat(prompt, temperature=0.25, max_tokens=320)

        response = jsonify({"success": True, "answer": answer})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"ask-question error: {e}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "healthy", "message": "StudyMate API is running"})


@app.route("/analyze-face", methods=["POST", "OPTIONS"])
def analyze_face_basic():
    """Lightweight camera analysis for attention tracking used by the StudyMate frontend."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        frame_b64 = data.get("frame") or data.get("image")
        if not frame_b64:
            return jsonify({"error": "No frame provided"}), 400

        img = decode_image_b64(frame_b64)
        if img is None:
            return jsonify({"error": "Invalid frame data"}), 400

        result = analyze_frame_basic(img)
        response = jsonify(result)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"analyze-face error: {e}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500


@app.route("/generate-notes-doc", methods=["POST", "OPTIONS"])
def generate_notes_doc():
    """Create downloadable notes (docx when available, otherwise txt) for a module."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        title = data.get("moduleTitle", "Module")
        summary = data.get("summary", "")
        points = data.get("keyPoints", [])

        safe_title = re.sub(r"[^\w\-]+", "_", title).strip("_") or "module"

        if Document:
            doc = Document()
            doc.add_heading(f"Notes: {title}", level=1)
            if summary:
                doc.add_paragraph("Summary:")
                doc.add_paragraph(summary)
            if points:
                doc.add_paragraph("Key Points:")
                for p in points[:5]:
                    doc.add_paragraph(p, style='List Bullet')

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            hex_blob = buf.read().hex()
            file_name = f"Notes_{safe_title}.docx"
        else:
            content = f"Notes: {title}\n\nSummary:\n{summary}\n\nKey Points:\n" + "\n".join(f"- {p}" for p in points[:5])
            hex_blob = content.encode("utf-8").hex()
            file_name = f"Notes_{safe_title}.txt"

        response = jsonify({
            "file_blob": hex_blob,
            "file_name": file_name
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"generate-notes-doc error: {e}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500


@app.route("/generate-challenge", methods=["POST", "OPTIONS"])
def generate_challenge():
    """Return a deterministic practice lab challenge (frontend caches by module)."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        module_title = data.get("moduleTitle", "Module")

        title = f"Build a Mini Page: {module_title}"
        question = (
            "Create a simple web page with a heading, a styled button, and a script that shows an alert when the button is clicked. "
            "Use semantic HTML, add basic CSS, and vanilla JS for interactivity."
        )
        starting_code = {
            "html": "<h1>Your Title</h1>\n<button id=\"actionBtn\">Click Me</button>",
            "css": "body { font-family: sans-serif; padding: 1rem; }\nbutton { padding: .5rem 1rem; }",
            "js": "document.getElementById('actionBtn').addEventListener('click', () => alert('Hello!'));"
        }

        response = jsonify({
            "type": "web",
            "title": title,
            "question": question,
            "starting_code": starting_code,
            "solution": starting_code
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"generate-challenge error: {e}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500


@app.route("/get-hint", methods=["POST", "OPTIONS"])
def get_hint():
    """Provide lightweight hints for practice lab submissions."""
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response

    try:
        data = request.get_json() or {}
        user_code = data.get("user_code", {"html": "", "css": "", "js": ""})
        try_count = int(data.get("try_count", 1))

        hints = []

        html_code = user_code.get("html", "")
        if "<h1" not in html_code:
            hints.append("Add a main heading using <h1> to describe the page.")
        if "button" not in html_code:
            hints.append("Include a <button id=\"actionBtn\"> element to attach your click handler.")

        css_code = user_code.get("css", "")
        if "font-family" not in css_code:
            hints.append("Set a readable font in CSS, e.g., body { font-family: sans-serif; }.")
        if "padding" not in css_code:
            hints.append("Add padding to the button for a better click area.")

        js_code = user_code.get("js", "")
        if "addEventListener" not in js_code or "click" not in js_code:
            hints.append("Attach a click event listener to the button using addEventListener('click', ...).")

        if try_count <= 1 and not hints:
            hints.append("Try running your code, then iterate on styling and interaction.")

        hint_text = "\n".join(hints) if hints else "Looks great! Consider improving accessibility and responsive styles."
        response = jsonify({"hint": hint_text})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
    except Exception as e:
        print(f"get-hint error: {e}")
        import traceback
        traceback.print_exc()
        response = jsonify({"error": str(e)})
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/detect-faces", methods=["POST", "OPTIONS"])
def detect_faces():
    """
    Endpoint to detect and count faces in an image.
    Used for exam proctoring to ensure only 1 person is present.
    Expects JSON with 'image' field containing base64 encoded image.
    """
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
    
    try:
        data = request.get_json()
        
        if not data or 'image' not in data:
            return jsonify({
                "success": False,
                "error": "No image data provided"
            }), 400
        
        image_data = data.get('image')
        timestamp = data.get('timestamp', None)
        exam_id = data.get('exam_id', None)
        
        # Decode base64 image
        if isinstance(image_data, str):
            # Remove data URL prefix if present
            if ',' in image_data:
                image_data = image_data.split(',')[1]
            img_bytes = base64.b64decode(image_data)
            img = Image.open(BytesIO(img_bytes))
            img_array = np.array(img)
        else:
            img_array = image_data
        
        # Convert RGB to BGR for OpenCV compatibility
        if len(img_array.shape) == 3 and img_array.shape[2] == 3:
            frame_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        else:
            frame_bgr = img_array
        
        # Detect faces using DeepFace
        try:
            faces = DeepFace.extract_faces(
                frame_bgr,
                detector_backend='opencv',
                enforce_detection=False,
                align=True
            )
            
            face_count = len(faces) if faces else 0
            
            # Determine status
            if face_count == 0:
                status = "no_face"
                message = "No face detected. Please ensure your face is visible and well-lit."
                allowed = False
            elif face_count == 1:
                status = "valid"
                message = "Perfect! Exactly 1 person detected."
                allowed = True
            else:
                status = "multiple_faces"
                message = f"Multiple people detected ({face_count} faces). Please ensure only you are in frame."
                allowed = False
            
            result = {
                "success": True,
                "face_count": face_count,
                "status": status,
                "message": message,
                "allowed": allowed,
                "timestamp": timestamp,
                "exam_id": exam_id
            }
            
            print(f"👤 Face Detection: {face_count} face(s) detected - Status: {status}")
            
            response = jsonify(result)
            response.headers.add("Access-Control-Allow-Origin", "*")
            return response
            
        except Exception as face_error:
            print(f"Face detection error: {str(face_error)}")
            # If detection fails, assume no face detected
            result = {
                "success": True,
                "face_count": 0,
                "status": "no_face",
                "message": "Could not detect faces. Please check camera and lighting.",
                "allowed": False,
                "timestamp": timestamp,
                "exam_id": exam_id
            }
            response = jsonify(result)
            response.headers.add("Access-Control-Allow-Origin", "*")
            return response
        
    except Exception as e:
        print(f"Error in face detection endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        response = jsonify({
            "success": False,
            "error": str(e),
            "message": "Failed to process image for face detection"
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/analyze-mood", methods=["POST", "OPTIONS"])
def analyze_mood():
    """
    Endpoint to analyze learner's mood and engagement from webcam frame.
    Expects JSON with 'image' field containing base64 encoded image.
    """
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
    
    try:
        data = request.get_json()
        
        if not data or 'image' not in data:
            return jsonify({
                "success": False,
                "error": "No image data provided"
            }), 400
        
        image_data = data.get('image')
        timestamp = data.get('timestamp', None)
        
        print(f"📸 Analyzing learner mood at timestamp: {timestamp}")
        
        # Analyze mood
        result = analyze_learner_mood(image_data)
        
        if result.get('success'):
            result['timestamp'] = timestamp
            print(f"✓ Emotion: {result['dominant_emotion']}, Status: {result['engagement_status']}, Score: {result['engagement_score']}")
        else:
            print(f"❌ Analysis failed: {result.get('error')}")
        
        response = jsonify(result)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"Error in mood analysis endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        response = jsonify({
            "success": False,
            "error": str(e)
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

@app.route("/batch-analyze-mood", methods=["POST", "OPTIONS"])
def batch_analyze_mood():
    """
    Endpoint to analyze multiple frames for engagement tracking over time.
    Expects JSON with 'frames' array containing objects with 'image' and 'timestamp'.
    """
    if request.method == "OPTIONS":
        response = jsonify({"status": "ok"})
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type")
        response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
        return response
    
    try:
        data = request.get_json()
        
        if not data or 'frames' not in data:
            return jsonify({
                "success": False,
                "error": "No frames data provided"
            }), 400
        
        frames = data.get('frames', [])
        
        if not frames:
            return jsonify({
                "success": False,
                "error": "Empty frames array"
            }), 400
        
        print(f"📸 Analyzing {len(frames)} frames for mood tracking")
        
        results = []
        engagement_timeline = []
        
        for idx, frame_data in enumerate(frames):
            image = frame_data.get('image')
            timestamp = frame_data.get('timestamp', idx)
            
            if image:
                analysis = analyze_learner_mood(image)
                if analysis.get('success'):
                    analysis['timestamp'] = timestamp
                    results.append(analysis)
                    engagement_timeline.append({
                        'timestamp': timestamp,
                        'score': analysis['engagement_score'],
                        'status': analysis['engagement_status'],
                        'emotion': analysis['dominant_emotion']
                    })
        
        # Calculate overall statistics
        if results:
            avg_engagement = sum(r['engagement_score'] for r in results) / len(results)
            emotion_counts = {}
            for r in results:
                emotion = r['dominant_emotion']
                emotion_counts[emotion] = emotion_counts.get(emotion, 0) + 1
            
            most_common_emotion = max(emotion_counts.items(), key=lambda x: x[1])[0]
            
            summary = {
                "success": True,
                "total_frames_analyzed": len(results),
                "average_engagement_score": round(avg_engagement, 2),
                "most_common_emotion": most_common_emotion,
                "emotion_distribution": emotion_counts,
                "engagement_timeline": engagement_timeline,
                "detailed_results": results
            }
        else:
            summary = {
                "success": False,
                "error": "No frames could be analyzed successfully"
            }
        
        print(f"✓ Batch analysis complete: {len(results)} frames processed")
        
        response = jsonify(summary)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        print(f"Error in batch mood analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        response = jsonify({
            "success": False,
            "error": str(e)
        })
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response, 500

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎓 StudyMate API Server")
    print("="*60)
    print(f"Mistral Model: {MISTRAL_MODEL}")
    print(f"Server: http://127.0.0.1:5000")
    print("="*60 + "\n")
    app.run(debug=True, host='127.0.0.1', port=5000)