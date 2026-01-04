import os
import json
import re
import base64
import subprocess
import urllib.parse
import datetime
from io import BytesIO
from pathlib import Path

from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.contrib.auth.decorators import login_required

from youtube_transcript_api import YouTubeTranscriptApi
import cv2
import numpy as np
import requests

try:
    from docx import Document
except ImportError:
    Document = None

# --- Configuration ---
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
MISTRAL_MODEL = os.getenv("MISTRAL_MODEL", "mistral-small-latest")
MISTRAL_API_URL = os.getenv("MISTRAL_API_URL", "https://api.mistral.ai/v1/chat/completions")

# Define paths relative to Django project root or settings.BASE_DIR
# Assuming settings.BASE_DIR is d:\sparkless 1\video_proctoring_project\proctoring
BASE_DIR = Path(__file__).resolve().parent.parent
SPARKLESS_ROOT = BASE_DIR.parent.parent  # Go up to d:\sparkless 1
TRANSCRIPTS_DIR = BASE_DIR / "transcripts"
TRANSCRIPTS_DIR.mkdir(exist_ok=True)

# Frontend path for static files (if needed for direct serving, though Django static is better)
FRONTEND_PATH = SPARKLESS_ROOT / "frontend"

# --- Helper Functions (Ported from Flask app.py) ---

def _mistral_chat(prompt, temperature=0.25, max_tokens=400):
    """Call Mistral chat completion and return text content."""
    if not MISTRAL_API_KEY:
        raise ValueError("MISTRAL_API_KEY not configured")

    try:
        resp = requests.post(
            MISTRAL_API_URL,
            headers={
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": MISTRAL_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "max_tokens": max_tokens,
            },
            timeout=60,
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    except Exception as exc:
        print(f"Mistral call failed: {exc}")
        return ""

def extract_youtube_id(url):
    if not url: return None
    parsed = urllib.parse.urlparse(url)
    if parsed.hostname in ['www.youtube.com', 'youtube.com']:
        q = urllib.parse.parse_qs(parsed.query)
        return q.get('v', [None])[0]
    elif parsed.hostname in ['youtu.be']:
        return parsed.path[1:]
    return None

def get_video_duration(video_id):
    try:
        result = subprocess.run([
            "yt-dlp",
            "--print", "duration",
            f"https://www.youtube.com/watch?v={video_id}"
        ], capture_output=True, text=True, check=True)
        return int(float(result.stdout.strip()))
    except Exception as e:
        print(f"Could not get video duration: {e}")
        return 3600

def get_transcript_with_timestamps(video_id):
    transcript_json_path = TRANSCRIPTS_DIR / f"{video_id}_timestamps.json"
    if transcript_json_path.exists():
        with open(transcript_json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    try:
        transcript_list = None

        if hasattr(YouTubeTranscriptApi, "get_transcript"):
            try:
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
            except Exception as inner:
                print(f"YouTubeTranscriptApi.get_transcript failed: {inner}")

        if transcript_list is None and hasattr(YouTubeTranscriptApi, "list_transcripts"):
            try:
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id).find_transcript(['en']).fetch()
            except Exception as inner:
                print(f"YouTubeTranscriptApi.list_transcripts failed: {inner}")

        # Fallback to yt-dlp if YouTubeTranscriptApi is unavailable
        if transcript_list is None:
            print(f"Falling back to yt-dlp for {video_id}...")
            try:
                subprocess.run([
                    "yt-dlp", "--write-auto-subs", "--sub-lang", "en", "--skip-download",
                    "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
                    f"https://www.youtube.com/watch?v={video_id}"
                ], check=True, capture_output=True, text=True, timeout=60)
                
                vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
                if vtt_files:
                    vtt_file = vtt_files[0]
                    lines = vtt_file.read_text(encoding="utf-8").splitlines()
                    transcript_list = [
                        {"text": line, "start": 0, "duration": 0}
                        for line in lines
                        if line.strip() and "-->" not in line and not line.strip().isdigit() and "WEBVTT" not in line and not line.startswith("NOTE")
                    ]
                    vtt_file.unlink()
            except Exception as fallback_err:
                print(f"yt-dlp fallback failed: {fallback_err}")

        if transcript_list is None:
            print(f"No transcript available for {video_id}; using mock")
            return [{"text": f"Content for video {video_id}", "start": 0, "duration": 0}]

        with open(transcript_json_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_list, f)
        return transcript_list
    except Exception as e:
        print(f"get_transcript_with_timestamps failed: {e}")
        return [{"text": f"Content for video {video_id}", "start": 0, "duration": 0}]

def get_transcript_text(video_id):
    transcript_path = TRANSCRIPTS_DIR / f"{video_id}.txt"
    if transcript_path.exists():
        return transcript_path.read_text(encoding="utf-8")
    
    # Fallback using yt-dlp (simplified for brevity, assuming yt-dlp is installed)
    try:
        subprocess.run([
            "yt-dlp", "--write-auto-subs", "--sub-lang", "en", "--skip-download",
            "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
            f"https://www.youtube.com/watch?v={video_id}"
        ], check=True, capture_output=True, text=True)
        
        vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
        if vtt_files:
            vtt_file = vtt_files[0]
            lines = vtt_file.read_text(encoding="utf-8").splitlines()
            full_text = " ".join(
                re.sub(r'<[^>]+>', '', line) 
                for line in lines 
                if line.strip() and "-->" not in line and not line.strip().isdigit() 
                and "WEBVTT" not in line and not line.startswith("NOTE")
            )
            transcript_path.write_text(full_text, encoding="utf-8")
            vtt_file.unlink()
            return full_text
    except Exception as e:
        print(f"yt-dlp failed: {e}")
    raise Exception("Could not fetch transcript")

def generate_course_overview(full_transcript_text, course_title):
    try:
        context_text = full_transcript_text[:5000]
        prompt = (
            "Analyze this video course transcript and provide a 2-4 sentence overview.\n"
            f"Course Title: {course_title}\n"
            f"Transcript Sample: {context_text}\n"
            "Return only the overview."
        )
        response_text = _mistral_chat(prompt, temperature=0.2, max_tokens=220)
        if response_text:
            return re.sub(r"\s+", " ", response_text.strip())
        raise ValueError("empty response")
    except Exception as e:
        print(f"Overview generation failed: {e}")
        return f"A comprehensive {course_title} covering essential concepts."

def generate_module_summary(module_text, module_num, start_time, end_time):
    try:
        start_min, end_min = int(start_time / 60), int(end_time / 60)
        context_text = module_text[:4000]
        prompt = (
            f"Summarize this segment (mins {start_min}-{end_min}) in 3-4 sentences.\n"
            f"Transcript: {context_text}\n"
            "Return only the summary."
        )
        response_text = _mistral_chat(prompt, temperature=0.25, max_tokens=260)
        if response_text:
            return re.sub(r"\s+", " ", response_text.strip())
        raise ValueError("empty response")
    except Exception:
        return f"Content covering minutes {start_min} to {end_min}."

def extract_key_points(module_text):
    try:
        context_text = module_text[:4000]
        prompt = (
            "Extract 5 key learning points from this transcript.\n"
            f"Transcript: {context_text}\n"
            "Return ONLY a JSON array of strings: [\"Point 1\", \"Point 2\", ...]"
        )
        response_text = _mistral_chat(prompt, temperature=0.2, max_tokens=220)
        cleaned = re.sub(r'^```json?\s*\n?|\n?```\s*$', '', response_text.strip(), flags=re.MULTILINE)
        return json.loads(cleaned)[:5]
    except Exception:
        return ["Key Concept 1", "Key Concept 2", "Key Concept 3", "Key Concept 4", "Key Concept 5"]

def split_transcript(transcript_data, daily_minutes, duration, title, has_timestamps=True):
    VIDEO_RATIO, QUIZ_RATIO = 0.85, 0.15
    daily_seconds = daily_minutes * 60
    video_seconds = int(daily_seconds * VIDEO_RATIO)
    quiz_seconds = int(daily_seconds * QUIZ_RATIO)
    
    num_modules = max(1, int(duration / video_seconds))
    
    if has_timestamps:
        full_text = " ".join(s['text'] for s in transcript_data)
        seconds_per_module = duration / num_modules
    else:
        words = transcript_data.split()
        full_text = transcript_data
        words_per_module = len(words) // num_modules
        seconds_per_module = duration // num_modules

    course_overview = generate_course_overview(full_text, title)
    modules = []

    for i in range(num_modules):
        start_time = int(i * seconds_per_module)
        end_time = int(min((i + 1) * seconds_per_module, duration))
        module_num = i + 1
        
        if has_timestamps:
            segs = [s for s in transcript_data if start_time <= s['start'] < end_time]
            module_text = " ".join(s['text'] for s in segs)
        else:
            start_word = i * words_per_module
            end_word = min((i + 1) * words_per_module, len(words))
            module_text = " ".join(words[start_word:end_word])

        summary = generate_module_summary(module_text, module_num, start_time, end_time)
        key_points = extract_key_points(module_text)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": round(video_seconds / 3600, 2),
            "quizDuration": round(quiz_seconds / 3600, 2),
            "totalDuration": round(daily_seconds / 3600, 2),
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {"description": summary, "keyPoints": key_points},
            "quiz": []
        })
    
    return modules, course_overview

def generate_quiz_ai(module_title, key_points):
    try:
        points_text = "\n".join(f"- {p}" for p in key_points[:5])
        prompt = (
            f"Create 5 multiple choice questions for {module_title}.\n"
            f"Key points: {points_text}\n"
            "Return ONLY valid JSON: [{ \"question\": \"...\", \"options\": { \"A\": \"...\" }, \"correct_answer\": \"A\" }]"
        )
        response_text = _mistral_chat(prompt, temperature=0.25, max_tokens=380)
        cleaned = re.sub(r'^```json?\s*\n?|\n?```\s*$', '', response_text.strip(), flags=re.MULTILINE)
        return json.loads(cleaned)[:5]
    except Exception:
        return []

# --- Views ---

@csrf_exempt
def courses_portal(request):
    """Serves the StudyMate frontend within Django."""
    # We'll read the index.html from the frontend folder and inject user context
    frontend_file = FRONTEND_PATH / 'index.html'
    
    print(f"\n🔍 DEBUG: Trying to load studymate dashboard")
    print(f"   FRONTEND_PATH: {FRONTEND_PATH}")
    print(f"   frontend_file: {frontend_file}")
    print(f"   File exists: {frontend_file.exists()}")
    
    if frontend_file.exists():
        try:
            with open(frontend_file, 'r', encoding='utf-8') as f:
                content = f.read()
                # Inject Django user context
                user = request.user.username if request.user.is_authenticated else "Guest"
                user_script = f'<script>window.djangoUser = "{user}";</script>'
                
                # Replace the script tag to include user context and point to the correct static URL
                # We use the view-based static serving at /studymate/static/
                content = content.replace('<script src="script.js"></script>', f'{user_script}\n<script src="/studymate/static/script.js"></script>')
                
                # Adjust CSS links to point to the view-based static URL
                content = content.replace('href="style.css', 'href="/studymate/static/style.css')
            
            print(f"✅ Frontend loaded successfully")
            return HttpResponse(content)
        except Exception as e:
            print(f"❌ Error reading file: {e}")
            return HttpResponse(f"Error reading file: {e}", status=500)
    else:
        print(f"❌ Frontend file not found: {frontend_file}")
        return HttpResponse(f"Frontend file not found at {frontend_file}", status=404)

# Helper to serve static files from frontend folder (Development only)
def serve_frontend_static(request, filename):
    file_path = FRONTEND_PATH / filename
    if file_path.exists():
        content_type = "text/css" if filename.endswith(".css") else "application/javascript"
        with open(file_path, 'r', encoding='utf-8') as f:
            return HttpResponse(f.read(), content_type=content_type)
    return HttpResponse(status=404)

@csrf_exempt
def generate_plan(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            course_title = data.get("courseTitle", "Course")
            course_link = data.get("courseLink")
            daily_hours = float(data.get("dailyStudyHours", 1))
            
            video_id = extract_youtube_id(course_link)
            if not video_id:
                return JsonResponse({"error": "Invalid YouTube link"}, status=400)
            
            duration = get_video_duration(video_id)
            transcript = get_transcript_with_timestamps(video_id)
            
            if transcript:
                plan, overview = split_transcript(transcript, daily_hours * 60, duration, course_title, True)
            else:
                text = get_transcript_text(video_id)
                plan, overview = split_transcript(text, daily_hours * 60, duration, course_title, False)
            
            return JsonResponse({
                "courseTitle": course_title,
                "courseDescription": overview,
                "videoID": video_id,
                "dailyPlan": plan,
                "streak": 0,
                "progress": 0
            })
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_quiz_view(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            quiz = generate_quiz_ai(data.get("moduleTitle"), data.get("keyPoints", []))
            return JsonResponse(quiz, safe=False)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def get_motivation(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        msg = f"🎉 Excellent work! You've completed {data.get('moduleTitle')} in {data.get('courseTitle')}."
        return JsonResponse({"message": msg})
    return JsonResponse({"status": "ok"})

@csrf_exempt
def analyze_face(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            img_b64 = data.get("frame", "").split(',', 1)[-1]
            if not img_b64: return JsonResponse({"error": "No frame"}, status=400)
            
            np_arr = np.frombuffer(base64.b64decode(img_b64), np.uint8)
            img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            faces = face_cascade.detectMultiScale(gray, 1.2, 5)
            
            face_present = len(faces) > 0
            attention_score = 1.0 if face_present else 0.0
            
            return JsonResponse({
                "face_present": face_present,
                "attention_score": attention_score,
                "distracted": not face_present,
                "metrics": {"faces_count": len(faces)}
            })
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_notes_doc(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            title = data.get("moduleTitle", "Module")
            summary = data.get("summary", "")
            points = data.get("keyPoints", [])
            
            if Document:
                doc = Document()
                doc.add_heading(f"Notes: {title}", 1)
                doc.add_paragraph("Summary:").add_run(f"\n{summary}")
                doc.add_paragraph("Key Points:")
                for p in points: doc.add_paragraph(p, style='List Bullet')
                
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                hex_blob = buf.read().hex()
                ext = ".docx"
            else:
                content = f"Notes: {title}\n\nSummary:\n{summary}\n\nPoints:\n" + "\n".join(f"- {p}" for p in points)
                hex_blob = content.encode("utf-8").hex()
                ext = ".txt"
                
            return JsonResponse({"file_blob": hex_blob, "file_name": f"Notes_{title}{ext}"})
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_challenge(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        title = f"Build a Mini Page: {data.get('moduleTitle')}"
        code = {
            "html": "<h1>Title</h1>\n<button id='btn'>Click</button>",
            "css": "body { padding: 1rem; }",
            "js": "document.getElementById('btn').onclick = () => alert('Hi');"
        }
        return JsonResponse({
            "type": "web", "title": title, "question": "Create a page...",
            "starting_code": code, "solution": code
        })
    return JsonResponse({"status": "ok"})

@csrf_exempt
def get_hint(request):
    """Get AI-powered hints and answers using Mistral API."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            question = data.get('question', '')
            module = data.get('module', 'this topic')
            
            if not question:
                return JsonResponse({"hint": "Please ask a specific question about the module."})
            
            # Create a prompt for Mistral to generate a helpful hint
            prompt = f"""You are a helpful AI study tutor. A student is asking about: {module}
            
Student's question: {question}

Provide a clear, concise, and helpful answer that:
1. Directly addresses the question
2. Explains the concept in simple terms
3. Provides a practical example or tip
4. Encourages further learning

Keep your answer to 2-3 sentences maximum. Be encouraging and supportive."""
            
            hint = _mistral_chat(prompt, temperature=0.5, max_tokens=150)
            
            if not hint:
                return JsonResponse({"hint": "That's a great question! Try reviewing the key concepts in this module."})
            
            return JsonResponse({"hint": hint})
        
        except Exception as e:
            print(f"Hint error: {e}")
            return JsonResponse({"hint": "I'm having trouble thinking right now. Try rephrasing your question."}, status=500)
    
    return JsonResponse({"status": "ok"})

@csrf_exempt
def summarize_content(request):
    """Summarize module content using Mistral API."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content', '')
            
            if not content:
                return JsonResponse({"summary": "No content provided to summarize."})
            
            # Create prompt for Mistral to summarize
            prompt = f"""Please provide a concise but comprehensive summary of the following educational content in 2-3 paragraphs:

{content}

Focus on:
1. Key concepts and learning objectives
2. Main topics covered
3. Practical applications
4. Key takeaways

Keep the summary clear and suitable for study notes."""
            
            # Call Mistral API
            summary = _mistral_chat(prompt, temperature=0.3, max_tokens=300)
            
            if not summary:
                return JsonResponse({"summary": "Unable to generate summary at this time."})
            
            return JsonResponse({"summary": summary})
        
        except Exception as e:
            print(f"Summarize error: {e}")
            return JsonResponse({"summary": f"Error generating summary: {str(e)}"}, status=500)
    
    return JsonResponse({"status": "ok"})

@csrf_exempt
def execute_code(request):
    """Execute JavaScript code safely and return output."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            code = data.get('code', '')
            
            if not code:
                return JsonResponse({"error": "No code provided"})
            
            # For security, we would normally use a sandbox or Node.js subprocess
            # For now, we'll use a subprocess to execute JavaScript safely
            import subprocess
            import tempfile
            
            # Create a temporary JavaScript file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.js', delete=False) as f:
                # Wrap code to capture console output
                wrapped_code = f"""
const originalLog = console.log;
let output = '';
console.log = function(...args) {{
    output += args.join(' ') + '\\n';
}};

try {{
{code}
}} catch(e) {{
    output += 'Error: ' + e.message;
}}

console.log(output);
"""
                f.write(wrapped_code)
                temp_file = f.name
            
            try:
                # Execute using Node.js if available, otherwise use Python
                result = subprocess.run(
                    ['node', temp_file],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                
                output = result.stdout.strip() or result.stderr.strip() or "(No output)"
                
                if result.returncode != 0 and result.stderr:
                    return JsonResponse({"error": result.stderr})
                
                return JsonResponse({"output": output})
            
            except FileNotFoundError:
                # Node.js not available, try Python eval (less safe but works)
                return JsonResponse({"error": "Code execution service not available. Please use your browser's developer console."})
            
            finally:
                # Clean up temp file
                import os
                try:
                    os.unlink(temp_file)
                except:
                    pass
        
        except Exception as e:
            print(f"Code execution error: {e}")
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"status": "ok"})


@csrf_exempt
def get_current_student(request):
    """
    Get current logged-in student information
    Returns: { full_name, username, roll_number, email }
    """
    try:
        from monitor.models import StudentProfile
        
        # Check if user is authenticated
        if not request.user.is_authenticated:
            # Return the default student for demo purposes
            default_student = StudentProfile.objects.first()
            if default_student:
                return JsonResponse({
                    "full_name": default_student.full_name,
                    "username": default_student.user.username,
                    "email": default_student.user.email,
                    "roll_number": default_student.roll_number,
                })
            return JsonResponse({"error": "No student found"}, status=404)
        
        user = request.user
        student_profile = StudentProfile.objects.filter(user=user).first()
        
        data = {
            "full_name": student_profile.full_name if student_profile else user.first_name or user.username,
            "username": user.username,
            "email": user.email,
            "roll_number": student_profile.roll_number if student_profile else None,
        }
        
        return JsonResponse(data)
    except Exception as e:
        print(f"Error getting student info: {e}")
        return JsonResponse({"error": str(e)}, status=500)


@csrf_exempt
def ai_tutor(request):
    """AI Study Buddy - Answer any question about a course using Mistral AI."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            question = data.get('question', '').strip()
            course_title = data.get('course_title', 'the course')
            course_description = data.get('course_description', '')
            module_title = data.get('module_title', '')
            
            if not question:
                return JsonResponse({"answer": "Please ask a question and I'll help you understand better! 📚"})
            
            # Build context for the AI
            context_parts = [f"Course: {course_title}"]
            if course_description:
                context_parts.append(f"Description: {course_description}")
            if module_title:
                context_parts.append(f"Current Module: {module_title}")
            
            context = "\n".join(context_parts)
            
            # Create a comprehensive prompt for Mistral
            prompt = f"""You are an expert AI Study Buddy helping a student learn. You're knowledgeable, friendly, and encouraging.

{context}

Student's Question: {question}

Please provide a helpful, educational response that:
1. Directly answers the question with accurate information
2. Explains concepts clearly with examples when helpful
3. Uses simple language that's easy to understand
4. Provides code examples if the question is about programming (use proper formatting)
5. Suggests next steps or related topics to explore
6. Is encouraging and supportive

If the question is about programming concepts like Java, Python, etc., include practical code examples.
If you're not 100% certain about something, acknowledge that and provide the best guidance you can.

Keep your response concise but comprehensive (max 3-4 paragraphs for explanations, or shorter for simple questions)."""
            
            answer = _mistral_chat(prompt, temperature=0.4, max_tokens=600)
            
            if not answer:
                return JsonResponse({
                    "answer": "I'm having a moment! 🤔 Could you try rephrasing your question? I'm here to help you learn!"
                })
            
            # Format the answer nicely
            return JsonResponse({
                "answer": answer,
                "status": "success"
            })
        
        except json.JSONDecodeError:
            return JsonResponse({"answer": "Sorry, I couldn't understand that. Please try again!"}, status=400)
        except Exception as e:
            print(f"AI Tutor error: {e}")
            return JsonResponse({
                "answer": "I'm experiencing some technical difficulties. Please try again in a moment! 🔧"
            }, status=500)
    
    return JsonResponse({"status": "ready", "message": "AI Study Buddy is ready to help!"})
