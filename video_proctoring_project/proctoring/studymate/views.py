import os
import json
import re
import base64
import subprocess
import urllib.parse
import datetime
from io import BytesIO
from pathlib import Path

from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.contrib.auth.decorators import login_required

import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
import cv2
import numpy as np

try:
    from docx import Document
except ImportError:
    Document = None

# --- Configuration ---
# In a real Django app, these should be in settings.py
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# Define paths relative to Django project root or settings.BASE_DIR
# Assuming settings.BASE_DIR is c:\sparkless\video_proctoring_project\proctoring
BASE_DIR = Path(__file__).resolve().parent.parent
TRANSCRIPTS_DIR = BASE_DIR / "transcripts"
TRANSCRIPTS_DIR.mkdir(exist_ok=True)

# Frontend path for static files (if needed for direct serving, though Django static is better)
FRONTEND_PATH = Path("c:/sparkless/frontend")

# --- Helper Functions (Ported from Flask app.py) ---

def extract_youtube_id(url):
    if not url: return None
    parsed = urllib.parse.urlparse(url)
    if parsed.hostname in ['www.youtube.com', 'youtube.com']:
        q = urllib.parse.parse_qs(parsed.query)
        return q.get('v', [None])[0]
    elif parsed.hostname in ['youtu.be']:
        return parsed.path[1:]
    return None

def get_video_duration(video_id):
    try:
        result = subprocess.run([
            "yt-dlp",
            "--print", "duration",
            f"https://www.youtube.com/watch?v={video_id}"
        ], capture_output=True, text=True, check=True)
        return int(float(result.stdout.strip()))
    except Exception as e:
        print(f"Could not get video duration: {e}")
        return 3600

def get_transcript_with_timestamps(video_id):
    transcript_json_path = TRANSCRIPTS_DIR / f"{video_id}_timestamps.json"
    if transcript_json_path.exists():
        with open(transcript_json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        with open(transcript_json_path, 'w', encoding='utf-8') as f:
            json.dump(transcript_list, f)
        return transcript_list
    except Exception as e:
        print(f"YouTubeTranscriptApi failed: {e}")
        return None

def get_transcript_text(video_id):
    transcript_path = TRANSCRIPTS_DIR / f"{video_id}.txt"
    if transcript_path.exists():
        return transcript_path.read_text(encoding="utf-8")
    
    # Fallback using yt-dlp (simplified for brevity, assuming yt-dlp is installed)
    try:
        subprocess.run([
            "yt-dlp", "--write-auto-subs", "--sub-lang", "en", "--skip-download",
            "-o", str(TRANSCRIPTS_DIR / f"{video_id}"),
            f"https://www.youtube.com/watch?v={video_id}"
        ], check=True, capture_output=True, text=True)
        
        vtt_files = list(TRANSCRIPTS_DIR.glob(f"{video_id}*.vtt"))
        if vtt_files:
            vtt_file = vtt_files[0]
            lines = vtt_file.read_text(encoding="utf-8").splitlines()
            full_text = " ".join(
                re.sub(r'<[^>]+>', '', line) 
                for line in lines 
                if line.strip() and "-->" not in line and not line.strip().isdigit() 
                and "WEBVTT" not in line and not line.startswith("NOTE")
            )
            transcript_path.write_text(full_text, encoding="utf-8")
            vtt_file.unlink()
            return full_text
    except Exception as e:
        print(f"yt-dlp failed: {e}")
    raise Exception("Could not fetch transcript")

def generate_course_overview(full_transcript_text, course_title):
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        context_text = full_transcript_text[:5000]
        prompt = f"""Analyze this video course transcript and provide a 2-4 sentence overview.
Course Title: {course_title}
Transcript Sample: {context_text}
Return only the overview."""
        response = model.generate_content(prompt)
        return re.sub(r'\s+', ' ', response.text.strip())
    except Exception as e:
        print(f"Overview generation failed: {e}")
        return f"A comprehensive {course_title} covering essential concepts."

def generate_module_summary(module_text, module_num, start_time, end_time):
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        start_min, end_min = int(start_time / 60), int(end_time / 60)
        context_text = module_text[:4000]
        prompt = f"""Summarize this segment (mins {start_min}-{end_min}) in 3-4 sentences.
Transcript: {context_text}
Return only the summary."""
        response = model.generate_content(prompt)
        return re.sub(r'\s+', ' ', response.text.strip())
    except Exception:
        return f"Content covering minutes {start_min} to {end_min}."

def extract_key_points(module_text):
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        context_text = module_text[:4000]
        prompt = f"""Extract 5 key learning points from this transcript.
Transcript: {context_text}
Return ONLY a JSON array of strings: ["Point 1", "Point 2", ...]"""
        response = model.generate_content(prompt)
        text = re.sub(r'^```json?\s*\n?|\n?```\s*$', '', response.text.strip(), flags=re.MULTILINE)
        return json.loads(text)[:5]
    except Exception:
        return ["Key Concept 1", "Key Concept 2", "Key Concept 3", "Key Concept 4", "Key Concept 5"]

def split_transcript(transcript_data, daily_minutes, duration, title, has_timestamps=True):
    VIDEO_RATIO, QUIZ_RATIO = 0.85, 0.15
    daily_seconds = daily_minutes * 60
    video_seconds = int(daily_seconds * VIDEO_RATIO)
    quiz_seconds = int(daily_seconds * QUIZ_RATIO)
    
    num_modules = max(1, int(duration / video_seconds))
    
    if has_timestamps:
        full_text = " ".join(s['text'] for s in transcript_data)
        seconds_per_module = duration / num_modules
    else:
        words = transcript_data.split()
        full_text = transcript_data
        words_per_module = len(words) // num_modules
        seconds_per_module = duration // num_modules

    course_overview = generate_course_overview(full_text, title)
    modules = []

    for i in range(num_modules):
        start_time = int(i * seconds_per_module)
        end_time = int(min((i + 1) * seconds_per_module, duration))
        module_num = i + 1
        
        if has_timestamps:
            segs = [s for s in transcript_data if start_time <= s['start'] < end_time]
            module_text = " ".join(s['text'] for s in segs)
        else:
            start_word = i * words_per_module
            end_word = min((i + 1) * words_per_module, len(words))
            module_text = " ".join(words[start_word:end_word])

        summary = generate_module_summary(module_text, module_num, start_time, end_time)
        key_points = extract_key_points(module_text)
        
        modules.append({
            "day": module_num,
            "title": f"Module {module_num}",
            "description": summary,
            "duration": round(video_seconds / 3600, 2),
            "quizDuration": round(quiz_seconds / 3600, 2),
            "totalDuration": round(daily_seconds / 3600, 2),
            "motivation": "Stay focused!",
            "completed": False,
            "startTime": start_time,
            "endTime": end_time,
            "module": {"description": summary, "keyPoints": key_points},
            "quiz": []
        })
    
    return modules, course_overview

def generate_quiz_ai(module_title, key_points):
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        points_text = "\n".join(f"- {p}" for p in key_points[:5])
        prompt = f"""Create 5 multiple choice questions for {module_title}.
Key points: {points_text}
Return ONLY valid JSON: [{{ "question": "...", "options": {{ "A": "...", ... }}, "correct_answer": "A" }}]"""
        response = model.generate_content(prompt)
        text = re.sub(r'^```json?\s*\n?|\n?```\s*$', '', response.text.strip(), flags=re.MULTILINE)
        return json.loads(text)[:5]
    except Exception:
        return []

# --- Views ---

@login_required
def courses_portal(request):
    """Serves the StudyMate frontend within Django."""
    # We'll read the index.html from the frontend folder and inject user context
    frontend_file = FRONTEND_PATH / 'index.html'
    if frontend_file.exists():
        with open(frontend_file, 'r', encoding='utf-8') as f:
            content = f.read()
            # Inject Django user context
            user_script = f'<script>window.djangoUser = "{request.user.username}";</script>'
            
            # Replace the script tag to include user context and point to the correct static URL
            # We use the view-based static serving at /studymate/static/
            content = content.replace('<script src="script.js"></script>', f'{user_script}\n<script src="/studymate/static/script.js"></script>')
            
            # Adjust CSS links to point to the view-based static URL
            content = content.replace('href="style.css', 'href="/studymate/static/style.css')
            
        return HttpResponse(content)
    else:
        return HttpResponse("Frontend file not found", status=404)

# Helper to serve static files from frontend folder (Development only)
def serve_frontend_static(request, filename):
    file_path = FRONTEND_PATH / filename
    if file_path.exists():
        content_type = "text/css" if filename.endswith(".css") else "application/javascript"
        with open(file_path, 'r', encoding='utf-8') as f:
            return HttpResponse(f.read(), content_type=content_type)
    return HttpResponse(status=404)

@csrf_exempt
def generate_plan(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            course_title = data.get("courseTitle", "Course")
            course_link = data.get("courseLink")
            daily_hours = float(data.get("dailyStudyHours", 1))
            
            video_id = extract_youtube_id(course_link)
            if not video_id:
                return JsonResponse({"error": "Invalid YouTube link"}, status=400)
            
            duration = get_video_duration(video_id)
            transcript = get_transcript_with_timestamps(video_id)
            
            if transcript:
                plan, overview = split_transcript(transcript, daily_hours * 60, duration, course_title, True)
            else:
                text = get_transcript_text(video_id)
                plan, overview = split_transcript(text, daily_hours * 60, duration, course_title, False)
            
            return JsonResponse({
                "courseTitle": course_title,
                "courseDescription": overview,
                "videoID": video_id,
                "dailyPlan": plan,
                "streak": 0,
                "progress": 0
            })
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_quiz_view(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            quiz = generate_quiz_ai(data.get("moduleTitle"), data.get("keyPoints", []))
            return JsonResponse(quiz, safe=False)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def get_motivation(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        msg = f"🎉 Excellent work! You've completed {data.get('moduleTitle')} in {data.get('courseTitle')}."
        return JsonResponse({"message": msg})
    return JsonResponse({"status": "ok"})

@csrf_exempt
def analyze_face(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            img_b64 = data.get("frame", "").split(',', 1)[-1]
            if not img_b64: return JsonResponse({"error": "No frame"}, status=400)
            
            np_arr = np.frombuffer(base64.b64decode(img_b64), np.uint8)
            img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            faces = face_cascade.detectMultiScale(gray, 1.2, 5)
            
            face_present = len(faces) > 0
            attention_score = 1.0 if face_present else 0.0
            
            return JsonResponse({
                "face_present": face_present,
                "attention_score": attention_score,
                "distracted": not face_present,
                "metrics": {"faces_count": len(faces)}
            })
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_notes_doc(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            title = data.get("moduleTitle", "Module")
            summary = data.get("summary", "")
            points = data.get("keyPoints", [])
            
            if Document:
                doc = Document()
                doc.add_heading(f"Notes: {title}", 1)
                doc.add_paragraph("Summary:").add_run(f"\n{summary}")
                doc.add_paragraph("Key Points:")
                for p in points: doc.add_paragraph(p, style='List Bullet')
                
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                hex_blob = buf.read().hex()
                ext = ".docx"
            else:
                content = f"Notes: {title}\n\nSummary:\n{summary}\n\nPoints:\n" + "\n".join(f"- {p}" for p in points)
                hex_blob = content.encode("utf-8").hex()
                ext = ".txt"
                
            return JsonResponse({"file_blob": hex_blob, "file_name": f"Notes_{title}{ext}"})
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"status": "ok"})

@csrf_exempt
def generate_challenge(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        title = f"Build a Mini Page: {data.get('moduleTitle')}"
        code = {
            "html": "<h1>Title</h1>\n<button id='btn'>Click</button>",
            "css": "body { padding: 1rem; }",
            "js": "document.getElementById('btn').onclick = () => alert('Hi');"
        }
        return JsonResponse({
            "type": "web", "title": title, "question": "Create a page...",
            "starting_code": code, "solution": code
        })
    return JsonResponse({"status": "ok"})

@csrf_exempt
def get_hint(request):
    if request.method == 'POST':
        return JsonResponse({"hint": "Try adding a header tag."})
    return JsonResponse({"status": "ok"})
